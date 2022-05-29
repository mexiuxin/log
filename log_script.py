#!/usr/local/python36/bin/python3.6
import os
from docx import Document
import random
import linecache
today=1
test=1
day1=1
month1=3
fileq=open('readfile','r')
ppp=len(fileq.readlines())
file = open('readfile','r')
### 输入你的名字
your_name="你的名字"
rand_file = open('rand_readfile','r')
file_line = len(rand_file.readlines()) 

while(today<=120):
    data = file.readline()
    # file = open("") 

    if(test==1):
        os.system("cp ../init.docx "+your_name+"实习日志第"+str(today)+"天.docx")
        doc = Document(your_name+"实习日志第"+str(today)+"天.docx")
        tables = doc.tables
        table0 = tables[0]
        cell_02=table0.cell(0, 2)
        cell_02.text=your_name

        cell_02=table0.cell(0, 4)
        cell_02.text="电子信息工程"

        cell_02=table0.cell(1, 2)
        cell_02.text="计算机网络"

        cell_02=table0.cell(1, 4)
        cell_02.text="19级"

        cell_02=table0.cell(2, 2)
        if(day1<=31 and month1==3):
            cell_02.text=str(month1) + " 月 " + str(day1) + " 日"
            if(day1==31 and month1==3):
                month1=month1+1
                day1=1
        elif(day1 <= 30 and month1 == 4):
            cell_02.text=str(month1) + " 月 " + str(day1) + " 日"
            if(day1==30 and month1 == 4):
                month1=month1+1
                day1=1
        elif(day1 <= 31 and month1 == 5):
            cell_02.text=str(month1) + " 月 " + str(day1) + " 日"


                

        cell_02=table0.cell(2, 4)
        cell_02.text="地区"

        cell_02=table0.cell(3, 2)
        cell_02.text="公司"

        cell_02=table0.cell(3, 4)
        cell_02.text="指导老师"
        cell_02=table0.cell(4, 2)
        cell_02.text="上班摸鱼工作给人跑腿"

        cell_02=table0.cell(5, 2)
        if(today == 1):
            cell_02.text="今天至于我来说，具有特殊的意义;因为今天是我在晨源劳务公司实习的第一天，也是我为不久的将来迈向社会大舞台的第一步，心里紧张而又激动。古人云天行健以自强不息，地势坤以厚德载物，对未来的实习生活，我愿意付出极大的热诚和努力。"
        elif(today <=ppp):
            cell_02.text=data
        else:
            rand=random.randint(1, file_line)
            cell_02.text=linecache.getline('rand_readfile', rand)



            
        #paragraphs = doc.paragraphs
        #par0 = paragraphs[2]
        # 获取段落文字
        #par0_string = par0.text
        #print(par0_string)

        paragraphs = doc.paragraphs

        i=0
        par0 = paragraphs[2]
        par0.text="第"+str(today)+"天"

        doc.save(your_name+"实习日志第"+str(today)+"天.docx")

    today=today+1
    day1=day1+1





